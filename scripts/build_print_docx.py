from __future__ import annotations

import datetime as dt
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_fonts(run, font_zh: str = "宋体", font_west: str = "Arial", size_pt: float | None = None) -> None:
    run.font.name = font_west
    if size_pt:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_zh)
    rfonts.set(qn("w:ascii"), font_west)
    rfonts.set(qn("w:hAnsi"), font_west)
    rfonts.set(qn("w:cs"), font_west)


def set_style_font(style, font_zh: str, font_west: str, size_pt: float | None = None) -> None:
    style.font.name = font_west
    if size_pt:
        style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_zh)
    rfonts.set(qn("w:ascii"), font_west)
    rfonts.set(qn("w:hAnsi"), font_west)
    rfonts.set(qn("w:cs"), font_west)


def add_field(paragraph, instruction: str, cached_text: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = cached_text
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, text_run, end])


def configure(doc: Document, margin_cm: float, font_zh: str, font_west: str, min_font_pt: float) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(margin_cm)
    sec.different_first_page_header_footer = True
    for name in ["Normal", "Body Text", "Header", "Footer"]:
        if name in doc.styles:
            set_style_font(doc.styles[name], font_zh, font_west, min_font_pt if name in ["Normal", "Body Text"] else 10.5)
    for name, size in [("Title", 24), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12.5)]:
        if name in doc.styles:
            set_style_font(doc.styles[name], font_zh, font_west, size)


def add_header_footer(doc: Document, title: str, font_zh: str, font_west: str) -> None:
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_fonts(header.add_run(title), font_zh, font_west, 10.5)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_fonts(footer.add_run("第 "), font_zh, font_west, 10.5)
    add_field(footer, "PAGE", "1")
    set_run_fonts(footer.add_run(" 页"), font_zh, font_west, 10.5)


def add_cover(doc: Document, title: str, subtitle: str, company: str, version: str, purpose: str, cover_image: Path | None, font_zh: str, font_west: str) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    set_run_fonts(run, font_zh, font_west, 24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_fonts(p.add_run(subtitle), font_zh, font_west, 14)
    if cover_image:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(cover_image), width=Cm(7.5))
        except Exception:
            pass
    for key, value in [
        ("公司", company),
        ("版本", version),
        ("日期", f"{dt.date.today().year} 年 {dt.date.today().month} 月"),
        ("用途", purpose),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_fonts(p.add_run(f"{key}：{value}"), font_zh, font_west, 11)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int, font_zh: str, font_west: str) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.bold = True
        set_run_fonts(run, font_zh, font_west)


def add_toc(doc: Document, sections: list[dict], font_zh: str, font_west: str, min_font_pt: float) -> None:
    add_heading(doc, "目录", 1, font_zh, font_west)
    for sec in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_fonts(p.add_run(sec["title"]), font_zh, font_west, min_font_pt)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, font_zh: str, font_west: str, min_font_pt: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    set_run_fonts(p.add_run(text), font_zh, font_west, min_font_pt)


def add_table(doc: Document, rows: list[list[str]], font_zh: str, font_west: str, min_font_pt: float) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_values in rows:
        cells = table.add_row().cells
        for i in range(cols):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[i].text = row_values[i] if i < len(row_values) else ""
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_run_fonts(r, font_zh, font_west, max(10.5, min_font_pt - 1))
    doc.add_paragraph("")


def add_album(doc: Document, images: list[Path], subject: str, figure_prefix: str, font_zh: str, font_west: str) -> None:
    for start in range(0, len(images), 2):
        pair = images[start : start + 2]
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col, img in enumerate(pair):
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(str(img), width=Cm(7.1))
            except Exception:
                set_run_fonts(p.add_run("图片无法插入"), font_zh, font_west, 10.5)
            cap = table.cell(1, col).paragraphs[0]
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_fonts(cap.add_run(f"图 {figure_prefix}-{start + col + 1} {subject}参考图（{start + col + 1}）"), font_zh, font_west, 10.5)
        if start + 2 < len(images):
            doc.add_page_break()


def build_print_docx(
    output_path: str | Path,
    title: str,
    sections: list[dict],
    images: list[Path],
    subtitle: str = "打印版",
    company: str = "",
    version: str = "V1.0",
    purpose: str = "正式打印 / 审阅 / 交付",
    figure_subject: str = "主题",
    figure_prefix: str = "8",
    font_zh: str = "宋体",
    font_west: str = "Arial",
    min_font_pt: float = 12,
    margin_cm: float = 2,
) -> Path:
    output = Path(output_path)
    doc = Document()
    configure(doc, margin_cm, font_zh, font_west, min_font_pt)
    add_header_footer(doc, title, font_zh, font_west)
    add_cover(doc, title, subtitle, company, version, purpose, images[0] if images else None, font_zh, font_west)
    add_toc(doc, sections, font_zh, font_west, min_font_pt)
    for sec in sections:
        add_heading(doc, sec["title"], 1, font_zh, font_west)
        for paragraph in sec.get("paragraphs", []):
            add_paragraph(doc, paragraph, font_zh, font_west, min_font_pt)
        for table in sec.get("tables", []):
            add_table(doc, table, font_zh, font_west, min_font_pt)
        if sec.get("album") and images:
            add_album(doc, images, figure_subject, figure_prefix, font_zh, font_west)
    doc.save(output)
    return output
