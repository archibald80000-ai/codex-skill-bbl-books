from __future__ import annotations

import re
from pathlib import Path

FORBIDDEN_TERMS = [
    "测试目录", "清洗报告", "处理报告", "素材清单", "原始素材", "版本差异", "相似度",
    "PackageNotFoundError", "error", "~$", "width=", "height=", "orientation=", "paragraphs=",
    "tables=", "embedded_images=", "AI废话", "删除多少条", "归并多少条", "skill", "脚本",
    "manifest", "本次处理目标", "本手册以测试目录", "请使用浏览器或 PDF 工具生成目录页",
]


def docx_text(path: Path) -> tuple[str, int, dict]:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    sec = doc.sections[0]
    rfonts = doc.styles["Normal"].element.rPr.rFonts
    meta = {
        "margin_2cm": all(round(x.cm, 2) == 2.0 for x in [sec.top_margin, sec.bottom_margin, sec.left_margin, sec.right_margin]),
        "font_zh": rfonts.get(qn("w:eastAsia")),
        "font_west": rfonts.get(qn("w:ascii")),
        "font_size": doc.styles["Normal"].font.size.pt if doc.styles["Normal"].font.size else None,
    }
    return "\n".join(chunks), len(doc.inline_shapes), meta


def validate_docx(path: Path, expected_images: int = 0) -> dict[str, bool | str]:
    text, image_count, meta = docx_text(path)
    return {
        "docx可打开": True,
        "无内部处理词": not any(term in text for term in FORBIDDEN_TERMS),
        "无图片文件名尺寸": not bool(re.search(r"[a-f0-9]{24,}\.(jpg|jpeg|png|webp)|width=|height=|orientation=", text, re.I)),
        "图片已纳入": image_count >= expected_images,
        "页边距2cm": bool(meta["margin_2cm"]),
        "中文宋体": meta["font_zh"] == "宋体",
        "西文Arial": meta["font_west"] == "Arial",
        "正文12pt": (meta["font_size"] or 0) >= 12,
    }


def validate_pdf(path: Path) -> dict[str, bool | str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    return {
        "pdf可读": path.exists() and path.stat().st_size > 0 and len(reader.pages) > 0,
        "pdf无内部处理词": not any(term in text for term in FORBIDDEN_TERMS),
        "pdf有页码": "第 2 页" in text or "第 3 页" in text,
    }

