from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from PIL import Image
from docx import Document

try:
    import openpyxl
except Exception:  # pragma: no cover
    openpyxl = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None


SUPPORTED_TEXT = {".md", ".txt", ".html", ".htm"}
SUPPORTED_IMAGES = {".jpg", ".jpeg", ".png", ".webp"}
SUPPORTED_DOCS = {".docx", ".pdf", ".xlsx", ".xls", ".csv"}
IGNORED_SUFFIXES = {".tmp", ".bak", ".log"}


@dataclass
class InputItem:
    path: Path
    kind: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


def should_ignore(path: Path) -> bool:
    name = path.name
    lower = name.lower()
    for part in path.parts:
        part_lower = part.lower()
        if part_lower.startswith("bbl-books_") or "成书版" in part or "打印版" in part:
            return True
    if name.startswith("~$") or name.startswith("."):
        return True
    if "__MACOSX" in path.parts:
        return True
    if path.suffix.lower() in IGNORED_SUFFIXES:
        return True
    if lower.startswith("bbl-books_") or lower.startswith("bbl-books"):
        return True
    if "成书版" in name or "打印版_v" in name:
        return True
    if any(token in lower for token in ["manifest", "素材清单", "清洗报告", "处理报告", "内部处理报告", "内部清洗报告"]):
        return True
    return False


def expand_inputs(inputs: list[str]) -> list[Path]:
    paths: list[Path] = []
    for raw in inputs:
        p = Path(raw).expanduser()
        if p.is_dir():
            for child in sorted(p.rglob("*")):
                if child.is_file() and not should_ignore(child):
                    paths.append(child)
        elif p.is_file() and not should_ignore(p):
            paths.append(p)
    seen: set[Path] = set()
    unique: list[Path] = []
    for path in paths:
        resolved = path.resolve()
        if resolved not in seen:
            seen.add(resolved)
            unique.append(resolved)
    return unique


def read_text(path: Path) -> list[str]:
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return [line.strip() for line in path.read_text(encoding=enc).splitlines() if line.strip()]
        except UnicodeDecodeError:
            continue
    return [path.read_text(errors="ignore").strip()]


def read_docx(path: Path) -> InputItem:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    return InputItem(path=path, kind="docx", paragraphs=paragraphs, tables=tables)


def read_pdf(path: Path) -> InputItem:
    paragraphs: list[str] = []
    if PdfReader:
        reader = PdfReader(str(path))
        for page in reader.pages:
            paragraphs.extend([line.strip() for line in (page.extract_text() or "").splitlines() if line.strip()])
    return InputItem(path=path, kind="pdf", paragraphs=paragraphs)


def read_xlsx(path: Path) -> InputItem:
    if openpyxl is None:
        return InputItem(path=path, kind="xlsx", metadata={"note": "openpyxl unavailable"})
    wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
    tables: list[list[list[str]]] = []
    paragraphs: list[str] = []
    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                rows.append(values)
        if rows:
            paragraphs.append(f"工作表：{ws.title}")
            tables.append(rows)
    return InputItem(path=path, kind="xlsx", paragraphs=paragraphs, tables=tables)


def read_csv(path: Path) -> InputItem:
    rows: list[list[str]] = []
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            with path.open("r", encoding=enc, newline="") as fh:
                rows = [[cell.strip() for cell in row] for row in csv.reader(fh)]
            break
        except UnicodeDecodeError:
            continue
    rows = [row for row in rows if any(row)]
    return InputItem(path=path, kind="csv", tables=[rows] if rows else [])


def inspect_image(path: Path) -> InputItem:
    with Image.open(path) as img:
        width, height = img.size
    return InputItem(path=path, kind="image", metadata={"width": width, "height": height})


def inspect_inputs(inputs: list[str]) -> list[InputItem]:
    items: list[InputItem] = []
    for path in expand_inputs(inputs):
        ext = path.suffix.lower()
        try:
            if ext == ".docx":
                items.append(read_docx(path))
            elif ext in SUPPORTED_TEXT:
                items.append(InputItem(path=path, kind=ext.lstrip("."), paragraphs=read_text(path)))
            elif ext == ".pdf":
                items.append(read_pdf(path))
            elif ext in {".xlsx", ".xls"}:
                items.append(read_xlsx(path))
            elif ext == ".csv":
                items.append(read_csv(path))
            elif ext in SUPPORTED_IMAGES:
                items.append(inspect_image(path))
        except Exception as exc:
            items.append(InputItem(path=path, kind="unreadable", metadata={"reason": f"{type(exc).__name__}: {exc}"}))
    return items
